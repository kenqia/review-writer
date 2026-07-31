from __future__ import annotations

import base64
import hashlib
import io
import json
from pathlib import Path
from unittest.mock import patch
from urllib.parse import quote

import pytest

from review_writer.project.source_truth import canonical_digest


TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)
HUMAN_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAIAAACQd1PeAAAADElEQVR4nGP4z8AAAAMBAQDJ/pLvAAAAAElFTkSuQmCC"
)
WORKFLOW_DIGEST = "a" * 64
LINEAGE_DIGEST = "b" * 64
SOURCE_ATTRIBUTION = (
    "Source Figure Attribution: source-figure-1 | source-a | page 3 | "
    "Figure 1"
)


def _write_json(path: Path, value: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(value, ensure_ascii=False, allow_nan=False, indent=2, sort_keys=True)
        + "\n",
        encoding="utf-8",
    )


def _chemical_lineage(*, with_dependency: bool = False) -> dict[str, object]:
    return {
        "chemical_paper_import_digests": [
            {
                "study_id": "study-a",
                "import_digest": "d" * 64,
                "state_digest": "e" * 64,
            }
        ],
        "chemical_paper_safe_summary": {
            "schema_version": "chemical-paper-safe-summary.v2",
            "route": "chemical-paper-zip-only",
            "study_count": 1,
            "molecule_count": 125,
            "missing_name_count": 0,
            "missing_resolved_smiles_count": 32,
            "ai_authored_smiles_count": 0,
            "element_review_counts": {
                "not_reviewed": 125,
                "confirmed": 0,
                "corrected": 0,
                "not_applicable": 0,
            },
            "reaction_data_status": "unavailable_not_provided",
        },
        "chemical_paper_claim_dependencies": (
            [
                {
                    "claim_id": "claim-a",
                    "study_id": "study-a",
                    "molecule_index": 0,
                    "required_fields": ["resolved_smiles"],
                    "requires_element_review": False,
                    "requires_reaction_data": False,
                }
            ]
            if with_dependency
            else []
        ),
    }


def _attach_chemical_lineage(
    project: Path, *, with_dependency: bool = False
) -> dict[str, object]:
    chemical = _chemical_lineage(with_dependency=with_dependency)
    lineage_path = project / "04_manuscript/manuscript_lineage.v2.json"
    lineage = json.loads(lineage_path.read_text(encoding="utf-8"))
    lineage.update(chemical)
    _write_json(lineage_path, lineage)
    return chemical


def _dependency_currentness(*, blocked: bool) -> dict[str, object]:
    reasons = ["CHEMICAL_REQUIRED_FIELD_UNRESOLVED"] if blocked else []
    return {
        "schema_version": "chemical-paper-dependency-currentness.v1",
        "lineage_binding_status": "current",
        "claims": [
            {
                "claim_id": "claim-a",
                "status": "needs_review" if blocked else "current",
                "dependencies": [
                    {
                        "study_id": "study-a",
                        "molecule_index": 0,
                        "status": "needs_review" if blocked else "current",
                        "required_field_statuses": {
                            "resolved_smiles": "unresolved" if blocked else "resolved"
                        },
                        "element_review_state": "not_reviewed",
                        "reaction_data_status": "unavailable_not_provided",
                        "blocking_reasons": reasons,
                    }
                ],
                "blocking_reasons": reasons,
            }
        ],
        "can_release": not blocked,
        "blocking_reasons": reasons,
    }


def _placeholder(*, status: str = "awaiting_human_figure") -> dict[str, object]:
    return {
        "placeholder_id": "synthesis-figure-1",
        "scientific_question": "How do the reported operating windows compare?",
        "reader_takeaway": "The studies use non-equivalent reporting windows.",
        "panels": [
            {
                "panel": "A",
                "task": "Compare the reported windows without ranking performance.",
                "synthesis_claim_ids": ["synthesis-1"],
                "source_figure_ids": ["source-figure-1"],
            }
        ],
        "comparison_axis": "reported operating window",
        "required_labels_units": ["temperature (C)"],
        "counter_evidence": ["Photon flux was not normalized."],
        "forbidden_overclaims": ["Do not rank manufacturing readiness."],
        "unresolved_uncertainties": ["Cross-study productivity is not comparable."],
        "caption_draft": "Reported operating windows; NR denotes not reported.",
        "target_size": "full width",
        "status": status,
    }


@pytest.fixture
def new_route_project(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    from review_writer.delivery import project_release

    project = tmp_path / "review-projects" / "new-route"
    (project / "01_evidence/source_truth/study-a").mkdir(parents=True)
    image = project / "03_figures/source.png"
    image.parent.mkdir(parents=True, exist_ok=True)
    image.write_bytes(TINY_PNG)
    manuscript = (
        "# Current evidence review\n\n"
        "## Evidence synthesis\n\n"
        "The approved evidence supports a bounded comparison. [synthesis:synthesis-1]\n\n"
        "![Reaction scope](../03_figures/source.png)\n\n"
        "Figure 1. Reaction scope.\n\n"
        f"{SOURCE_ATTRIBUTION}\n\n"
        "SYNTHESIS_FIGURE_PLACEHOLDER: synthesis-figure-1\n\n"
        "Scientific question: How do the reported operating windows compare?\n\n"
        "Figure task: Compare the reported windows without ranking performance.\n\n"
        "Status: awaiting human figure.\n\n"
        "## References\n\n"
        "[1] Synthetic reference entry.\n"
    )
    manuscript_path = project / "04_manuscript/manuscript.md"
    manuscript_path.parent.mkdir(parents=True)
    manuscript_path.write_text(manuscript, encoding="utf-8")
    _write_json(
        project / "04_manuscript/manuscript_lineage.v2.json",
        {
            "schema_version": "manuscript-lineage.v2",
            "route": "evidence-to-release.v1",
            "manuscript_sha256": hashlib.sha256(manuscript.encode()).hexdigest(),
            "lineage_digest": LINEAGE_DIGEST,
            "source_figure_registry_digest": canonical_digest([]),
            "synthesis_figure_placeholder_digest": canonical_digest([_placeholder()]),
        },
    )
    source_figure = {
        "figure_id": "source-figure-1",
        "study_id": "study-a",
        "source_id": "source-a",
        "page": 3,
        "figure_label": "Figure 1",
        "caption": "Reaction scope.",
        "asset_path": "03_figures/source.png",
        "asset_sha256": hashlib.sha256(TINY_PNG).hexdigest(),
        "source_pdf_sha256": "c" * 64,
        "evidence_ids": ["evidence-1"],
        "selection_status": "selected",
        "selection_reason": "Supports the evidence synthesis section.",
    }
    _write_json(
        project / "03_figures/source_figure_registry.json",
        {
            "schema_version": "review-writer-source-figure-registry.v1",
            "project_id": project.name,
            "figures": [source_figure],
            "registry_digest": canonical_digest([source_figure]),
        },
    )
    _write_json(
        project / "03_figures/synthesis_figure_placeholders.json",
        {"placeholders": [_placeholder()]},
    )
    legacy = project / "04_first_draft/first_draft.md"
    legacy.parent.mkdir(parents=True)
    legacy.write_text("# LEGACY SENTINEL\n", encoding="utf-8")

    workflow = {
        "route": "evidence-to-release.v1",
        "parse_ready": True,
        "paper_evidence_ready": True,
        "synthesis_ready": True,
        "section_contracts_ready": True,
        "manuscript_ready": True,
        "internal_draft_export_ready": True,
        "verified_release_ready": False,
        "workflow_digest": WORKFLOW_DIGEST,
    }
    monkeypatch.setattr(project_release, "workflow_state", lambda _: dict(workflow))
    monkeypatch.setattr(
        project_release,
        "manuscript_state",
        lambda _: {
            "workflow_can_continue": True,
            "reason_code": "MANUSCRIPT_APPROVED",
            "manuscript_sha256": hashlib.sha256(manuscript.encode()).hexdigest(),
            "lineage_digest": LINEAGE_DIGEST,
        },
        raising=False,
    )
    return project


def _release_bytes(project: Path) -> dict[str, tuple[bytes, int]]:
    stage = project / "05_release"
    return {
        path.name: (path.read_bytes(), path.stat().st_mtime_ns)
        for path in stage.glob("*")
        if path.is_file()
    }


def _dashboard_request(dashboard, review_root: Path, raw_request: bytes) -> tuple[int, dict, bytes]:
    class FakeSocket:
        def __init__(self, incoming: bytes) -> None:
            self.input = io.BytesIO(incoming)
            self.output = io.BytesIO()

        def makefile(self, mode: str, *args, **kwargs):
            return self.input if "r" in mode else self.output

        def sendall(self, data: bytes) -> None:
            self.output.write(data)

        def close(self) -> None:
            pass

    dashboard.DashboardHandler.review_root = review_root
    socket = FakeSocket(raw_request)
    dashboard.DashboardHandler(socket, ("127.0.0.1", 0), object())
    head, body = socket.output.getvalue().split(b"\r\n\r\n", 1)
    lines = head.decode("iso-8859-1").split("\r\n")
    headers = dict(line.split(": ", 1) for line in lines[1:] if ": " in line)
    return int(lines[0].split()[1]), headers, body


def test_new_route_internal_release_reads_only_manuscript_v2_and_writes_release_artifacts(
    new_route_project: Path,
) -> None:
    from review_writer.delivery.project_release import build_project_release

    result = build_project_release(
        new_route_project, release_level="SELF_REVIEWED_DRAFT"
    )

    assert result["status"] == "SELF_REVIEWED_DRAFT"
    assert result["release_level"] == "SELF_REVIEWED_DRAFT"
    assert result["placeholder_count"] == 1
    assert Path(result["snapshot"]).relative_to(new_route_project).as_posix() == (
        "05_release/self_reviewed_draft.md"
    )
    assert Path(result["docx"]).relative_to(new_route_project).as_posix() == (
        "05_release/self_reviewed_draft.docx"
    )
    assert "LEGACY SENTINEL" not in Path(result["snapshot"]).read_text(encoding="utf-8")
    snapshot = json.loads(
        (new_route_project / "05_release/release_snapshot.json").read_text(encoding="utf-8")
    )
    assert snapshot["workflow_digest"] == WORKFLOW_DIGEST
    assert snapshot["lineage_digest"] == LINEAGE_DIGEST
    assert snapshot["integrity"]["markdown_roundtrip_match"] is True
    assert snapshot["integrity"]["attribution_complete"] is True
    assert snapshot["integrity"]["provenance_valid"] is True

    from docx import Document

    document = Document(result["docx"])
    assert document.core_properties.title == "new-route - SELF_REVIEWED_DRAFT"
    assert document.core_properties.subject == "review-writer project new-route"
    assert document.core_properties.author == "review-writer"
    assert document.core_properties.last_modified_by == "review-writer"
    assert document.core_properties.keywords == (
        "new-route; SELF_REVIEWED_DRAFT; review-writer"
    )


def test_internal_docx_binds_chemical_lineage_and_adds_explicit_limitations(
    new_route_project: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    from docx import Document
    from review_writer.delivery.project_release import build_project_release

    chemical = _attach_chemical_lineage(new_route_project)
    from review_writer.delivery import project_release

    monkeypatch.setattr(
        project_release,
        "dependency_currentness_for_project",
        lambda *_args, **_kwargs: {
            "schema_version": "chemical-paper-dependency-currentness.v1",
            "lineage_binding_status": "current",
            "claims": [],
            "can_release": True,
            "blocking_reasons": [],
        },
    )

    result = build_project_release(
        new_route_project, release_level="SELF_REVIEWED_DRAFT"
    )

    released_markdown = Path(result["snapshot"]).read_text(encoding="utf-8")
    assert "Original PDFs remain the scientific source of truth" in released_markdown
    assert "does not mean zero confirmed reactions" in released_markdown
    document_text = "\n".join(
        paragraph.text for paragraph in Document(result["docx"]).paragraphs
    )
    assert "Chemical Paper output is a manual-export parsing aid, not scientific truth" in document_text
    snapshot = json.loads(
        (new_route_project / "05_release/release_snapshot.json").read_text(encoding="utf-8")
    )
    assert snapshot["chemical_paper_binding_digest"] == canonical_digest(chemical)
    assert snapshot["chemical_paper_safe_summary"]["missing_resolved_smiles_count"] == 32
    assert "import_digest" not in json.dumps(snapshot["chemical_paper_safe_summary"])


def test_dual_parse_release_is_bound_and_stale_state_is_zero_write(
    new_route_project: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from review_writer.delivery import project_release
    from review_writer.delivery.project_release import (
        ProjectReleaseError,
        build_project_release,
    )

    (new_route_project / "01_evidence/dual_source").mkdir()
    lineage_path = new_route_project / "04_manuscript/manuscript_lineage.v2.json"
    lineage = json.loads(lineage_path.read_text(encoding="utf-8"))
    lineage["dual_parse_bindings"] = [
        {
            "study_id": "study-a",
            "source_tier": "core",
            "requires_chemical": True,
            "dual_source_binding_digest": "1" * 64,
            "generic_version": "2" * 64,
            "chemical_version": "3" * 64,
            "chemical_completion_digest": "4" * 64,
            "reconciliation_digest": "5" * 64,
        }
    ]
    _write_json(lineage_path, lineage)
    current = {
        "dual_parse_status": "current",
        "internal_release_ready": True,
        "expert_release_ready": False,
        "hard_fails": [],
        "issues": ["CHEMICAL_REACTION_DATA_UNAVAILABLE", "SYNTHESIS_FIGURE_PENDING"],
        "reaction_data_status": "unavailable_not_provided",
        "reaction_count": None,
        "credits_status": "NOT_APPLICABLE_BY_CURRENT_SCOPE",
    }
    monkeypatch.setattr(
        project_release, "dual_parse_release_state", lambda _: dict(current)
    )

    result = build_project_release(
        new_route_project, release_level="SELF_REVIEWED_DRAFT"
    )

    binding_digest = canonical_digest(lineage["dual_parse_bindings"])
    assert result["dual_parse_status"] == "current"
    assert result["dual_parse_binding_digest"] == binding_digest
    quality = json.loads(
        (new_route_project / "05_release/quality_report.json").read_text(
            encoding="utf-8"
        )
    )
    assert quality["dual_parse_status"] == "current"
    assert quality["dual_parse_binding_digest"] == binding_digest
    assert quality["reaction_data_status"] == "unavailable_not_provided"
    assert quality["reaction_count"] is None
    assert quality["credits_status"] == "NOT_APPLICABLE_BY_CURRENT_SCOPE"
    assert "credits" not in quality
    docx = Path(result["docx"])
    assert project_release.new_route_release_docx_is_current(docx) is True

    before = _release_bytes(new_route_project)
    current.update(
        {
            "dual_parse_status": "stale",
            "internal_release_ready": False,
            "hard_fails": ["CORE_GENERIC_PARSE_MISSING_OR_STALE"],
        }
    )
    with pytest.raises(
        ProjectReleaseError, match="CORE_GENERIC_PARSE_MISSING_OR_STALE"
    ):
        build_project_release(
            new_route_project, release_level="SELF_REVIEWED_DRAFT"
        )
    assert _release_bytes(new_route_project) == before
    assert project_release.new_route_release_docx_is_current(docx) is False


def test_expert_release_fails_closed_on_used_unresolved_chemical_field_without_overwrite(
    new_route_project: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    from review_writer.delivery import project_release
    from review_writer.delivery.project_release import ProjectReleaseError, build_project_release

    _attach_chemical_lineage(
        new_route_project,
        with_dependency=True,
    )
    monkeypatch.setattr(
        project_release,
        "dependency_currentness_for_project",
        lambda *_args, **_kwargs: _dependency_currentness(blocked=True),
    )
    stage = new_route_project / "05_release"
    stage.mkdir(parents=True)
    for name in (
        "expert_reviewed_release.md",
        "expert_reviewed_release.docx",
        "release_snapshot.json",
        "quality_report.json",
    ):
        (stage / name).write_bytes(f"sentinel:{name}".encode())
    before = _release_bytes(new_route_project)

    with pytest.raises(ProjectReleaseError, match="CHEMICAL_DEPENDENCY_UNRESOLVED"):
        build_project_release(
            new_route_project, release_level="EXPERT_REVIEWED_RELEASE"
        )

    assert _release_bytes(new_route_project) == before


def test_internal_docx_allows_explicitly_dependent_gap_but_discloses_release_limit(
    new_route_project: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    from review_writer.delivery import project_release
    from review_writer.delivery.project_release import build_project_release

    _attach_chemical_lineage(new_route_project, with_dependency=True)
    monkeypatch.setattr(
        project_release,
        "dependency_currentness_for_project",
        lambda *_args, **_kwargs: _dependency_currentness(blocked=True),
    )

    result = build_project_release(
        new_route_project, release_level="SELF_REVIEWED_DRAFT"
    )

    released_markdown = Path(result["snapshot"]).read_text(encoding="utf-8")
    assert "claim-dependent chemical values still require review" in released_markdown
    snapshot = json.loads(
        (new_route_project / "05_release/release_snapshot.json").read_text(
            encoding="utf-8"
        )
    )
    assert snapshot["chemical_paper_dependency_can_release"] is False


def test_released_docx_becomes_stale_when_chemical_binding_changes(
    new_route_project: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    from review_writer.delivery import project_release
    from review_writer.delivery.project_release import (
        build_project_release,
        new_route_release_docx_is_current,
    )

    _attach_chemical_lineage(new_route_project)
    monkeypatch.setattr(
        project_release,
        "dependency_currentness_for_project",
        lambda *_args, **_kwargs: {
            "schema_version": "chemical-paper-dependency-currentness.v1",
            "lineage_binding_status": "current",
            "claims": [],
            "can_release": True,
            "blocking_reasons": [],
        },
    )
    result = build_project_release(
        new_route_project, release_level="SELF_REVIEWED_DRAFT"
    )
    assert new_route_release_docx_is_current(Path(result["docx"])) is True

    lineage_path = new_route_project / "04_manuscript/manuscript_lineage.v2.json"
    lineage = json.loads(lineage_path.read_text(encoding="utf-8"))
    lineage["chemical_paper_safe_summary"]["missing_resolved_smiles_count"] = 31
    _write_json(lineage_path, lineage)

    assert new_route_release_docx_is_current(Path(result["docx"])) is False


def test_expert_release_rejects_pending_placeholder_without_overwrite(
    new_route_project: Path,
) -> None:
    from review_writer.delivery.project_release import ProjectReleaseError, build_project_release

    stage = new_route_project / "05_release"
    stage.mkdir(parents=True)
    for name in (
        "expert_reviewed_release.md",
        "expert_reviewed_release.docx",
        "release_snapshot.json",
        "quality_report.json",
    ):
        (stage / name).write_bytes(f"sentinel:{name}".encode())
    before = _release_bytes(new_route_project)

    with pytest.raises(ProjectReleaseError, match="FIGURE_PLACEHOLDER_PENDING"):
        build_project_release(
            new_route_project, release_level="EXPERT_REVIEWED_RELEASE"
        )

    assert _release_bytes(new_route_project) == before


def test_expert_release_accepts_rights_cleared_human_verified_synthesis_figure(
    new_route_project: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from review_writer.delivery import project_release
    from review_writer.delivery.project_release import build_project_release

    registry_path = new_route_project / "03_figures/source_figure_registry.json"
    registry = json.loads(registry_path.read_text(encoding="utf-8"))
    registry["figures"][0].update(
        {
            "rights_status": "cleared",
            "rights_license": "Publisher permission recorded",
        }
    )
    _write_json(registry_path, registry)

    placeholder = _placeholder(status="verified")
    placeholders = [placeholder]
    _write_json(
        new_route_project / "03_figures/synthesis_figure_placeholders.json",
        {"placeholders": placeholders},
    )
    human_asset = new_route_project / "03_figures/synthesis-figure-1.png"
    human_asset.write_bytes(HUMAN_PNG)
    manuscript_path = new_route_project / "04_manuscript/manuscript.md"
    manuscript = manuscript_path.read_text(encoding="utf-8")
    manuscript = manuscript.replace(
        "SYNTHESIS_FIGURE_PLACEHOLDER: synthesis-figure-1\n\n"
        "Scientific question: How do the reported operating windows compare?\n\n"
        "Figure task: Compare the reported windows without ranking performance.\n\n"
        "Status: awaiting human figure.\n",
        "![Human verified synthesis](../03_figures/synthesis-figure-1.png)\n\n"
        "HUMAN_SYNTHESIS_FIGURE: 03_figures/synthesis-figure-1.png\n",
    )
    manuscript_path.write_text(manuscript, encoding="utf-8")
    manuscript_sha256 = hashlib.sha256(manuscript.encode()).hexdigest()
    lineage_path = new_route_project / "04_manuscript/manuscript_lineage.v2.json"
    lineage = json.loads(lineage_path.read_text(encoding="utf-8"))
    lineage["manuscript_sha256"] = manuscript_sha256
    lineage["synthesis_figure_placeholder_digest"] = canonical_digest(placeholders)
    _write_json(lineage_path, lineage)

    asset_sha256 = hashlib.sha256(HUMAN_PNG).hexdigest()
    verification_object = {
        "placeholder_digest": canonical_digest(placeholders),
        "placeholder_id": "synthesis-figure-1",
        "asset_path": "03_figures/synthesis-figure-1.png",
        "asset_sha256": asset_sha256,
        "lineage_digest": LINEAGE_DIGEST,
    }
    _write_json(
        new_route_project / "03_figures/synthesis_figure_verification.json",
        {
            "verifications": [
                {
                    **verification_object,
                    "verification": {
                        "schema_version": "verification-decision.v1",
                        "actor_type": "human_researcher",
                        "actor_label": "test-human",
                        "action": "verify",
                        "reason": "Verified against the approved scientific figure brief.",
                        "decided_at": "2026-07-29T00:00:00Z",
                        "bound_object_digest": canonical_digest(verification_object),
                        "bound_gate_digest": LINEAGE_DIGEST,
                    },
                }
            ]
        },
    )
    monkeypatch.setattr(
        project_release,
        "manuscript_state",
        lambda _: {
            "workflow_can_continue": True,
            "reason_code": "MANUSCRIPT_APPROVED",
            "manuscript_sha256": manuscript_sha256,
            "lineage_digest": LINEAGE_DIGEST,
        },
    )

    result = build_project_release(
        new_route_project, release_level="EXPERT_REVIEWED_RELEASE"
    )

    assert result["status"] == "EXPERT_REVIEWED_RELEASE"
    assert result["pending_placeholder_count"] == 0
    assert Path(result["docx"]).name == "expert_reviewed_release.docx"
    from view import serve_review_dashboard as dashboard

    assert dashboard.project_release_docx_is_current(Path(result["docx"])) is True


def test_failed_new_route_export_preserves_existing_release(
    new_route_project: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from review_writer.delivery import project_release
    from review_writer.delivery.project_release import ProjectReleaseError, build_project_release

    build_project_release(new_route_project, release_level="SELF_REVIEWED_DRAFT")
    before = _release_bytes(new_route_project)
    monkeypatch.setattr(
        project_release.subprocess,
        "run",
        lambda *args, **kwargs: project_release.subprocess.CompletedProcess(args[0], 1, "", "failed"),
    )

    with pytest.raises(ProjectReleaseError, match="DOCX_EXPORT_FAILED"):
        build_project_release(new_route_project, release_level="SELF_REVIEWED_DRAFT")

    assert _release_bytes(new_route_project) == before


def test_new_route_rejects_legacy_docx_repackage_and_preserves_current_release(
    new_route_project: Path,
) -> None:
    from review_writer.delivery.project_release import ProjectReleaseError, build_project_release

    first = build_project_release(new_route_project, release_level="SELF_REVIEWED_DRAFT")
    legacy_docx = new_route_project / "05_final_audit/final_draft.docx"
    legacy_docx.parent.mkdir(parents=True)
    legacy_docx.write_bytes(Path(first["docx"]).read_bytes())
    before = _release_bytes(new_route_project)

    with pytest.raises(ProjectReleaseError, match="LEGACY_REPACKAGE_ONLY"):
        build_project_release(new_route_project, release_level="SELF_REVIEWED_DRAFT")

    assert _release_bytes(new_route_project) == before


def test_new_route_release_currentness_fails_closed_after_workflow_drift(
    new_route_project: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from review_writer.delivery import project_release
    from review_writer.delivery.project_release import build_project_release
    from view import serve_review_dashboard as dashboard

    result = build_project_release(new_route_project, release_level="SELF_REVIEWED_DRAFT")
    docx = Path(result["docx"])
    assert dashboard.is_project_release_docx(docx, new_route_project.parents[1]) is True
    assert dashboard.project_release_docx_is_current(docx) is True
    ready = dashboard.project_final_payload(
        new_route_project.parents[1], new_route_project.name
    )
    assert ready["release_status"] == "SELF_REVIEWED_DRAFT"
    assert ready["manuscript_source"] == "release_snapshot"
    assert ready["final_draft_docx_exists"] is True

    monkeypatch.setattr(
        project_release,
        "workflow_state",
        lambda _: {
            "route": "evidence-to-release.v1",
            "parse_ready": True,
            "paper_evidence_ready": True,
            "synthesis_ready": True,
            "section_contracts_ready": True,
            "manuscript_ready": True,
            "internal_draft_export_ready": True,
            "verified_release_ready": False,
            "workflow_digest": "d" * 64,
        },
    )

    assert dashboard.project_release_docx_is_current(docx) is False
    stale = dashboard.project_final_payload(
        new_route_project.parents[1], new_route_project.name
    )
    assert stale["release_status"] == "RELEASE_OUTDATED"
    assert stale["final_draft_docx_exists"] is False
    request = (
        f"GET /file?path={quote(str(docx), safe='')} HTTP/1.1\r\n"
        "Host: localhost\r\n\r\n"
    ).encode("ascii")
    status, _, _ = _dashboard_request(
        dashboard, new_route_project.parents[1], request
    )
    assert status == 403


def test_new_route_export_api_contract_passes_release_level_as_strict_json(
    tmp_path: Path,
) -> None:
    from view import serve_review_dashboard as dashboard

    body = b'{"release_level":"EXPERT_REVIEWED_RELEASE"}'
    expected = {
        "ok": True,
        "release_level": "EXPERT_REVIEWED_RELEASE",
        "filename": "expert_reviewed_release.docx",
        "size": 123,
        "release_status": "EXPERT_REVIEWED_RELEASE",
    }
    request = (
        b"POST /api/project/project-a/export-docx HTTP/1.1\r\n"
        b"Host: localhost\r\nContent-Type: application/json\r\nContent-Length: "
        + str(len(body)).encode("ascii")
        + b"\r\n\r\n"
        + body
    )
    with patch.object(dashboard, "export_project_docx", return_value=expected) as export:
        status, headers, response = _dashboard_request(dashboard, tmp_path, request)

    assert status == 200
    assert headers["Content-Type"] == "application/json; charset=utf-8"
    assert json.loads(response) == expected
    export.assert_called_once_with(
        tmp_path, "project-a", release_level="EXPERT_REVIEWED_RELEASE"
    )


def test_current_release_download_has_unambiguous_attachment_filename(
    new_route_project: Path,
) -> None:
    from review_writer.delivery.project_release import build_project_release
    from view import serve_review_dashboard as dashboard

    result = build_project_release(
        new_route_project, release_level="SELF_REVIEWED_DRAFT"
    )
    docx = Path(result["docx"])
    request = (
        f"GET /file?path={quote(str(docx), safe='')} HTTP/1.1\r\n"
        "Host: localhost\r\n\r\n"
    ).encode("ascii")

    status, headers, _ = _dashboard_request(
        dashboard, new_route_project.parents[1], request
    )

    assert status == 200
    assert headers["Content-Disposition"] == (
        "attachment; filename*=UTF-8''self_reviewed_draft.docx"
    )


@pytest.mark.parametrize(
    "body",
    [
        b"{}",
        b'{"release_level":"UNSUPPORTED"}',
        b'{"release_level":"SELF_REVIEWED_DRAFT","extra":true}',
        b"not-json",
    ],
)
def test_new_route_export_api_rejects_invalid_contract_as_json(
    tmp_path: Path, body: bytes
) -> None:
    from view import serve_review_dashboard as dashboard

    request = (
        b"POST /api/project/project-a/export-docx HTTP/1.1\r\n"
        b"Host: localhost\r\nContent-Type: application/json\r\nContent-Length: "
        + str(len(body)).encode("ascii")
        + b"\r\n\r\n"
        + body
    )
    with patch.object(dashboard, "export_project_docx") as export:
        status, headers, response = _dashboard_request(dashboard, tmp_path, request)

    assert status == 400
    assert headers["Content-Type"] == "application/json; charset=utf-8"
    assert json.loads(response) == {
        "ok": False,
        "error_code": "RELEASE_REQUEST_INVALID",
        "message": "release request must contain exactly one supported release_level",
    }
    export.assert_not_called()


def test_new_route_export_api_maps_release_gate_to_json_conflict(tmp_path: Path) -> None:
    from review_writer.delivery.project_release import ProjectReleaseError
    from view import serve_review_dashboard as dashboard

    body = b'{"release_level":"EXPERT_REVIEWED_RELEASE"}'
    request = (
        b"POST /api/project/project-a/export-docx HTTP/1.1\r\n"
        b"Host: localhost\r\nContent-Type: application/json\r\nContent-Length: "
        + str(len(body)).encode("ascii")
        + b"\r\n\r\n"
        + body
    )
    with patch.object(
        dashboard,
        "export_project_docx",
        side_effect=ProjectReleaseError(
            "FIGURE_PLACEHOLDER_PENDING", "expert release requires verification"
        ),
    ):
        status, headers, response = _dashboard_request(dashboard, tmp_path, request)

    assert status == 409
    assert headers["Content-Type"] == "application/json; charset=utf-8"
    assert json.loads(response) == {
        "ok": False,
        "error_code": "FIGURE_PLACEHOLDER_PENDING",
        "message": "expert release requires verification",
    }


def test_new_route_export_api_hides_unexpected_backend_error(tmp_path: Path) -> None:
    from view import serve_review_dashboard as dashboard

    body = b'{"release_level":"SELF_REVIEWED_DRAFT"}'
    request = (
        b"POST /api/project/project-a/export-docx HTTP/1.1\r\n"
        b"Host: localhost\r\nContent-Type: application/json\r\nContent-Length: "
        + str(len(body)).encode("ascii")
        + b"\r\n\r\n"
        + body
    )
    with patch.object(
        dashboard, "export_project_docx", side_effect=RuntimeError("private detail")
    ):
        status, headers, response = _dashboard_request(dashboard, tmp_path, request)

    assert status == 500
    assert headers["Content-Type"] == "application/json; charset=utf-8"
    assert json.loads(response) == {
        "ok": False,
        "error_code": "RELEASE_INTERNAL_ERROR",
        "message": "release service failed",
    }


def test_explicit_internal_level_keeps_legacy_export_compatible(tmp_path: Path) -> None:
    from view import serve_review_dashboard as dashboard

    project = tmp_path / "review-projects/legacy"
    docx = project / "05_final_audit/final_draft.docx"
    docx.parent.mkdir(parents=True)
    docx.write_bytes(b"legacy-docx")
    with patch.object(
        dashboard,
        "build_project_release",
        return_value={"status": "AI_REVIEWED_BENCHMARK", "docx": docx},
    ):
        result = dashboard.export_project_docx(
            tmp_path, "legacy", release_level="SELF_REVIEWED_DRAFT"
        )

    assert result == {
        "ok": True,
        "release_level": "SELF_REVIEWED_DRAFT",
        "filename": "final_draft.docx",
        "size": len(b"legacy-docx"),
        "release_status": "AI_REVIEWED_BENCHMARK",
    }
