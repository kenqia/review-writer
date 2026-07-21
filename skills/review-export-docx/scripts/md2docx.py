#!/usr/bin/env python3
"""
md2docx.py -- Convert a Markdown file to DOCX using review_template.docx styles.

Inline support:
  **bold**  *italic*  ***bold-italic***  `code`
  ^superscript^       alnum_subscript_   $math$  $$display math$$
  [@citation]  ->  [citation key]

Section-aware styling (auto-detected from headings OR bold-only paragraphs):
  Abstract / Keywords / Acknowledgments / References / Supporting Information

Font specification (explicitly applied to every run):
  H1 title      : Times New Roman 18 pt  (xiao-er)
  Author line   : Times New Roman 12 pt  (xiao-si)
  Affiliation   : Times New Roman 10.5 pt (wu-hao)
  H2 heading    : Times New Roman 14 pt  (si-hao)  bold
  H3 heading    : Times New Roman 12 pt  bold italic
  Body / Abstract / Keywords : Times New Roman 12 pt
  Captions / References      : Times New Roman 10.5 pt

Usage:
    python3 scripts/md2docx.py --input review.md --output review.docx
"""
from __future__ import annotations

import argparse
import re
import sys
from copy import deepcopy  # noqa: F401
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml  # noqa: F401
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:
    from review_writer.docx_links import link_citations_and_dois
except ModuleNotFoundError:
    # Keep the repo-local skill executable when invoked by absolute path.
    _REPO_ROOT = Path(__file__).resolve().parents[3]
    if str(_REPO_ROOT) not in sys.path:
        sys.path.insert(0, str(_REPO_ROOT))
    from review_writer.docx_links import link_citations_and_dois

try:
    from latex2word import LatexToWordElement
    _LATEX_OK = True
except ImportError:
    _LATEX_OK = False

# ---------------------------------------------------------------------------
# Template style name map
# ---------------------------------------------------------------------------

_S: Dict[str, str] = {
    "title":        "BA_Title",
    "author":       "BB_Author_Name",
    "address":      "BC_Author_Address",
    "email":        "BI_Email_Address",
    "abstract":     "BD_Abstract",
    "keywords":     "BG_Keywords",
    "body":         "TA_Main_Text",
    "figure":       "VA_Figure_Caption",
    "table_title":  "VD_Table_Title",
    "table_body":   "TC_Table_Body",
    "chart":        "VB_Chart_Title",
    "scheme":       "VC_Scheme_Title",
    "references":   "TF_References_Section",
    "acks":         "TD_Acknowledgments",
    "supporting":   "TE_Supporting_Information",
    "footnote":     "FA_Corresponding_Author_Footnote",
}

# ---------------------------------------------------------------------------
# Font spec -- every run gets an explicit font name + size
# ---------------------------------------------------------------------------

_FONT_SPEC: Dict[str, Dict] = {
    "title":        {"font": "Times New Roman", "size": 18},
    "author":       {"font": "Times New Roman", "size": 12},
    "address":      {"font": "Times New Roman", "size": 10.5},
    "email":        {"font": "Times New Roman", "size": 10.5},
    "abstract":     {"font": "Times New Roman", "size": 12},
    "keywords":     {"font": "Times New Roman", "size": 12},
    "body":         {"font": "Times New Roman", "size": 12},
    "h2":           {"font": "Times New Roman", "size": 14,  "bold": True},
    "h3":           {"font": "Times New Roman", "size": 12,  "bold": True, "italic": True},
    "h4":           {"font": "Times New Roman", "size": 12,  "italic": True},
    "figure":       {"font": "Times New Roman", "size": 10.5},
    "table_title":  {"font": "Times New Roman", "size": 12},
    "table_body":   {"font": "Times New Roman", "size": 10.5},
    "scheme":       {"font": "Times New Roman", "size": 10.5},
    "chart":        {"font": "Times New Roman", "size": 10.5},
    "references":   {"font": "Times New Roman", "size": 10.5},
    "acks":         {"font": "Times New Roman", "size": 12},
    "supporting":   {"font": "Times New Roman", "size": 12},
    "footnote":     {"font": "Times New Roman", "size": 10.5},
}

# Heading level -> (para_style_key, font_spec_key)
_HEADING_FORMAT: Dict[int, Tuple[str, str]] = {
    1: ("title", "title"),
    2: ("body",  "h2"),
    3: ("body",  "h3"),
    4: ("body",  "h4"),
    5: ("body",  "body"),
    6: ("body",  "body"),
}

_SECTION_CONTEXT: Dict[str, str] = {
    "abstract":               "abstract",
    "keywords":               "keywords",
    "key words":              "keywords",
    "acknowledgments":        "acks",
    "acknowledgements":       "acks",
    "supporting information": "supporting",
    "references":             "references",
    "reference":              "references",
}

_CAPTION_PATTERNS: List[Tuple[re.Pattern, str]] = [
    (re.compile(r"^(figure|fig\.)\s*\d+", re.I), "figure"),
    (re.compile(r"^table\s*\d+",            re.I), "table_title"),
    (re.compile(r"^scheme\s*\d+",           re.I), "scheme"),
    (re.compile(r"^chart\s*\d+",            re.I), "chart"),
]


def _usable_page_width_inches(doc: Document) -> float:
    section = doc.sections[0]
    width_emu = section.page_width - section.left_margin - section.right_margin
    # 914400 EMUs per inch. Keep a conservative upper bound for journal templates.
    return max(1.0, min(6.2, width_emu / 914400))

_UNICODE_SUPERSCRIPT_MAP: Dict[str, str] = {
    "⁰": "0",
    "¹": "1",
    "²": "2",
    "³": "3",
    "⁴": "4",
    "⁵": "5",
    "⁶": "6",
    "⁷": "7",
    "⁸": "8",
    "⁹": "9",
    "⁺": "+",
    "⁻": "-",
    "⁼": "=",
    "⁽": "(",
    "⁾": ")",
}

_UNICODE_SUBSCRIPT_MAP: Dict[str, str] = {
    "₀": "0",
    "₁": "1",
    "₂": "2",
    "₃": "3",
    "₄": "4",
    "₅": "5",
    "₆": "6",
    "₇": "7",
    "₈": "8",
    "₉": "9",
    "₊": "+",
    "₋": "-",
    "₌": "=",
    "₍": "(",
    "₎": ")",
}

# ---------------------------------------------------------------------------
# Run dataclass
# ---------------------------------------------------------------------------

@dataclass
class Run:
    text:        str  = ""
    bold:        bool = False
    italic:      bool = False
    code:        bool = False
    superscript: bool = False
    subscript:   bool = False
    math:        str  = ""

# ---------------------------------------------------------------------------
# Inline parser
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r"(\$\$[\s\S]*?\$\$"
    r"|\$[^$\n]+?\$"
    r"|\*\*\*(?:\S[^*\n]*?\S|\S)\*\*\*"
    r"|\*\*(?:\S[^*\n]*?\S|\S)\*\*"
    r"|\^[^\^\s\n]+?\^"
    r"|_[^_\s\n]+_"
    r"|\*(?:\S[^*\n]*?\S|\S)\*"
    r"|__(?:\S[^_\n]*?\S|\S)__"
    r"|_(?:\S[^_\n]*?\S|\S)_"
    r"|`[^`\n]+`"
    r"|\[@([^\]]+)\]"
    r"|\[([^\]]*)\]\([^\)]*\)"
    r")"
)


def _parse_nested(inner: str, bold: bool = False, italic: bool = False) -> List[Run]:
    runs = parse_inline(inner)
    for r in runs:
        if bold:
            r.bold = True
        if italic:
            r.italic = True
    return runs


def parse_inline(raw: str) -> List[Run]:
    runs: List[Run] = []
    pos = 0
    for m in _INLINE_RE.finditer(raw):
        if m.start() > pos:
            runs.append(Run(text=raw[pos:m.start()]))
        token = m.group(0)
        char_before = raw[m.start() - 1] if m.start() > 0 else ""

        if token.startswith("$$"):
            runs.append(Run(math=token[2:-2].strip()))
        elif token.startswith("$"):
            runs.append(Run(math=token[1:-1].strip()))
        elif token.startswith("***"):
            runs.extend(_parse_nested(token[3:-3], bold=True, italic=True))
        elif token.startswith("**"):
            runs.extend(_parse_nested(token[2:-2], bold=True))
        elif token.startswith("^") and token.endswith("^"):
            runs.append(Run(text=token[1:-1], superscript=True))
        elif token.startswith("_") and token.endswith("_") and " " not in token[1:-1]:
            if char_before.isalnum():
                runs.append(Run(text=token[1:-1], subscript=True))
            else:
                runs.append(Run(text=token[1:-1], italic=True))
        elif token.startswith("*"):
            runs.extend(_parse_nested(token[1:-1], italic=True))
        elif token.startswith("__"):
            runs.extend(_parse_nested(token[2:-2], bold=True))
        elif token.startswith("_"):
            runs.extend(_parse_nested(token[1:-1], italic=True))
        elif token.startswith("`"):
            runs.append(Run(text=token[1:-1], code=True))
        elif token.startswith("[@"):
            cite_key = m.group(1) or token[2:-1]
            runs.append(Run(text=f"[{cite_key}]"))
        elif token.startswith("["):
            display = m.group(2)
            runs.append(Run(text=display if display is not None else token))
        else:
            runs.append(Run(text=token))
        pos = m.end()

    if pos < len(raw):
        runs.append(Run(text=raw[pos:]))
    return runs or [Run(text=raw)]


# ---------------------------------------------------------------------------
# Font + run application
# ---------------------------------------------------------------------------

def _apply_math(para, latex: str) -> None:
    if _LATEX_OK:
        try:
            LatexToWordElement(latex).add_latex_to_paragraph(para)
            return
        except Exception:
            pass
    run = para.add_run(f"[{latex}]")
    run.italic = True


def _split_script_segments(text: str) -> List[Tuple[str, str]]:
    segments: List[Tuple[str, str]] = []
    mode = "normal"
    current = ""

    def looks_like_ascii_subscript(index: int) -> bool:
        if index < 0 or index >= len(text):
            return False
        char = text[index]
        if not char.isdigit() or index == 0:
            return False
        prev = text[index - 1]
        if prev in "-–—/[ ":
            return False
        if prev.isalpha():
            return True
        if prev in ")]}" and index >= 2 and text[index - 2].isalpha():
            return True
        return False

    def flush() -> None:
        nonlocal current
        if current:
            segments.append((mode, current))
            current = ""

    for idx, char in enumerate(text):
        if char in _UNICODE_SUPERSCRIPT_MAP:
            char_mode = "superscript"
            rendered = _UNICODE_SUPERSCRIPT_MAP[char]
        elif char in _UNICODE_SUBSCRIPT_MAP:
            char_mode = "subscript"
            rendered = _UNICODE_SUBSCRIPT_MAP[char]
        elif looks_like_ascii_subscript(idx):
            char_mode = "subscript"
            rendered = char
        else:
            char_mode = "normal"
            rendered = char
        if char_mode != mode:
            flush()
            mode = char_mode
        current += rendered
    flush()
    return segments or [("normal", text)]


def apply_runs(
    para,
    runs: List[Run],
    spec_key: str = "body",
    force_bold: bool = False,
    force_italic: bool = False,
) -> None:
    spec        = _FONT_SPEC.get(spec_key, _FONT_SPEC["body"])
    font_name   = spec["font"]
    size_pt     = spec["size"]
    spec_bold   = spec.get("bold", False)
    spec_italic = spec.get("italic", False)

    for r in runs:
        if r.math:
            _apply_math(para, r.math)
            continue
        segments = [("normal", r.text)] if r.code else _split_script_segments(r.text)
        for segment_mode, segment_text in segments:
            if not segment_text:
                continue
            wr = para.add_run(segment_text)
            if r.code:
                wr.font.name = "Courier New"
                wr.font.size = Pt(9)
            else:
                wr.font.name = font_name
                wr.font.size = Pt(size_pt)
            wr.bold = (spec_bold or force_bold or r.bold) or None
            wr.italic = (spec_italic or force_italic or r.italic) or None
            if r.superscript or segment_mode == "superscript":
                wr.font.superscript = True
            if r.subscript or segment_mode == "subscript":
                wr.font.subscript = True


# ---------------------------------------------------------------------------
# Paragraph factory
# ---------------------------------------------------------------------------

def _para(
    doc: Document,
    style_key: str,
    spec_key: str,
    inline_text: str = "",
    force_bold: bool = False,
    force_italic: bool = False,
):
    p = doc.add_paragraph(style=_S.get(style_key, _S["body"]))
    if inline_text:
        apply_runs(p, parse_inline(inline_text),
                   spec_key=spec_key,
                   force_bold=force_bold,
                   force_italic=force_italic)
    return p


# ---------------------------------------------------------------------------
# Table builder
# ---------------------------------------------------------------------------

def _set_cell_borders(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")
        tcPr.append(elem)


def _add_table(doc: Document, header: List[str], rows: List[List[str]]) -> None:
    ncols = max(len(header), max((len(r) for r in rows), default=1))
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    for j, h in enumerate(header):
        cell = table.cell(0, j)
        cell.text = ""
        cell.paragraphs[0].style = doc.styles[_S["table_body"]]
        apply_runs(cell.paragraphs[0], parse_inline(h),
                   spec_key="table_body", force_bold=True)
        _set_cell_borders(cell)
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i + 1, j)
            cell.text = ""
            cell.paragraphs[0].style = doc.styles[_S["table_body"]]
            apply_runs(cell.paragraphs[0],
                       parse_inline(row[j] if j < len(row) else ""),
                       spec_key="table_body")
            _set_cell_borders(cell)


# ---------------------------------------------------------------------------
# Block tokenizer
# ---------------------------------------------------------------------------

@dataclass
class Block:
    kind:     str
    level:    int             = 0
    text:     str             = ""
    ordered:  bool            = False
    depth:    int             = 0
    code:     str             = ""
    language: str             = ""
    header:   List[str]       = field(default_factory=list)
    rows:     List[List[str]] = field(default_factory=list)
    alt:      str             = ""
    path:     str             = ""
    latex:    str             = ""
    lines:    List[str]       = field(default_factory=list)


_HEADING_RE    = re.compile(r"^(#{1,6})\s+(.*)")
_EMBEDDED_HEADING_PREFIX_RE = re.compile(r"^#{1,6}\s+")
_NUMBERED_SECTION_HEADING_RE = re.compile(r"^\d+(?:\.\d+)*\.\s+\S")
_HTML_ANCHOR_RE = re.compile(r"^<a\s+id=[\"']ref-\d+[\"']\s*>\s*</a>\s*$", re.I)
_UL_RE         = re.compile(r"^(\s*)[-*+]\s+(.*)")
_OL_RE         = re.compile(r"^(\s*)\d+[.)]\s+(.*)")
_FENCE_RE      = re.compile(r"^```(\w*)\s*$")
_MATH_FENCE_RE = re.compile(r"^\$\$\s*$")
_IMG_RE        = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
_TABLE_ROW_RE  = re.compile(r"^\|.+")
_HR_RE         = re.compile(r"^(?:-{3,}|_{3,}|\*{3,})\s*$")
_REF_ENTRY_RE  = re.compile(r"^\[?\d+\]?[.)\s]|\[@[^\]]+\]:")
_AFFIL_START   = re.compile(r"^\^[0-9,]+\^")
_INDENTED_RE   = re.compile(r"^(?: {4,}|\t+)(.*)$")


def _is_continuation(line: str) -> bool:
    if not line.strip():
        return False
    if _REF_ENTRY_RE.match(line):
        return False
    if _AFFIL_START.match(line):
        return False
    for pat in (_HEADING_RE, _FENCE_RE, _TABLE_ROW_RE,
                _UL_RE, _OL_RE, _HR_RE, _IMG_RE):
        if pat.match(line):
            return False
    return True


def tokenize(md_text: str) -> List[Block]:
    lines = md_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: List[Block] = []
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]

        # YAML front matter
        if i == 0 and line.strip() == "---":
            i += 1
            while i < n and lines[i].strip() != "---":
                i += 1
            i += 1
            continue

        # Fenced code block
        m = _FENCE_RE.match(line)
        if m:
            lang = m.group(1)
            code_lines: List[str] = []
            i += 1
            while i < n and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append(Block(kind="code_block", language=lang,
                                code="\n".join(code_lines)))
            continue

        # Display math fence
        if _MATH_FENCE_RE.match(line):
            math_lines: List[str] = []
            i += 1
            while i < n and not _MATH_FENCE_RE.match(lines[i]):
                math_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append(Block(kind="math_block", latex="\n".join(math_lines)))
            continue

        if _HTML_ANCHOR_RE.match(line.strip()):
            i += 1
            continue

        # ATX heading
        m = _HEADING_RE.match(line)
        if m:
            heading_text = m.group(2).strip()
            while _EMBEDDED_HEADING_PREFIX_RE.match(heading_text):
                heading_text = _EMBEDDED_HEADING_PREFIX_RE.sub("", heading_text, count=1).strip()
            blocks.append(Block(kind="heading", level=len(m.group(1)),
                                text=heading_text))
            i += 1
            continue

        # Horizontal rule
        if _HR_RE.match(line):
            blocks.append(Block(kind="hr"))
            i += 1
            continue

        # Standalone image
        m = _IMG_RE.match(line)
        if m:
            blocks.append(Block(kind="image", alt=m.group(1), path=m.group(2)))
            i += 1
            continue

        # Indented text block: preserve one source line -> one logical block line.
        m = _INDENTED_RE.match(line)
        if m and not _TABLE_ROW_RE.match(line):
            block_lines: List[str] = [m.group(1).rstrip()]
            i += 1
            while i < n:
                next_match = _INDENTED_RE.match(lines[i])
                if not next_match or not next_match.group(1).strip():
                    break
                block_lines.append(next_match.group(1).rstrip())
                i += 1
            blocks.append(Block(kind="indented_block", lines=block_lines))
            continue

        # Reference definition  [@key]: text...
        ref_m = re.match(r"^\[@([^\]]+)\]:\s*(.+)$", line)
        if ref_m:
            blocks.append(Block(kind="ref_def", text=ref_m.group(2).strip()))
            i += 1
            continue

        # Pipe table
        if _TABLE_ROW_RE.match(line):
            raw_rows: List[List[str]] = []
            while i < n and _TABLE_ROW_RE.match(lines[i]):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                raw_rows.append(cells)
                i += 1
            data = [r for r in raw_rows
                    if not all(re.match(r"^[-: ]+$", c) for c in r)]
            if data:
                blocks.append(Block(kind="table", header=data[0], rows=data[1:]))
            continue

        # Ordered list item
        m = _OL_RE.match(line)
        if m:
            blocks.append(Block(kind="list_item", ordered=True,
                                depth=len(m.group(1)) // 2, text=m.group(2)))
            i += 1
            continue

        # Unordered list item
        m = _UL_RE.match(line)
        if m:
            blocks.append(Block(kind="list_item", ordered=False,
                                depth=len(m.group(1)) // 2, text=m.group(2)))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Paragraph
        para_lines = [line]
        i += 1
        while i < n and _is_continuation(lines[i]):
            para_lines.append(lines[i].rstrip())
            i += 1
        blocks.append(Block(kind="paragraph",
                            text=" ".join(l.rstrip() for l in para_lines)))

    return blocks


# ---------------------------------------------------------------------------
# Context helpers
# ---------------------------------------------------------------------------

def _plain_text(raw: str) -> str:
    t = re.sub(r"\*+|__?", "", raw)
    t = re.sub(r"\[@[^\]]+\]", "", t)
    t = re.sub(r"\^[^\^]+\^", "", t)
    t = re.sub(r"`[^`]+`", "", t)
    t = re.sub(r"\[[^\]]*\]\([^\)]*\)", "", t)
    return t.strip()


def _section_ctx(text: str) -> Optional[str]:
    return _SECTION_CONTEXT.get(text.strip().lower())


def _caption_style(raw_text: str) -> Optional[str]:
    plain = _plain_text(raw_text)
    for pat, key in _CAPTION_PATTERNS:
        if pat.match(plain):
            return key
    return None


def _should_include_in_toc(text: str) -> bool:
    normalized = text.strip().lower()
    if normalized in {
        "table of contents",
        "abstract",
        "keywords",
        "key words",
        "acknowledgments",
        "acknowledgements",
        "references",
        "reference",
    }:
        return False
    return True


def _collect_static_toc_entries(blocks: List[Block]) -> List[Tuple[int, str]]:
    entries: List[Tuple[int, str]] = []
    for block in blocks:
        if block.kind != "heading":
            continue
        text = block.text.strip()
        effective_level = 2 if block.level == 1 and _NUMBERED_SECTION_HEADING_RE.match(text) else block.level
        if effective_level not in {2, 3, 4}:
            continue
        if not text or not _should_include_in_toc(text):
            continue
        entries.append((effective_level, text))
    return entries


def _insert_static_toc(doc: Document, entries: List[Tuple[int, str]]) -> None:
    for level, text in entries:
        p = doc.add_paragraph(style=_S["body"])
        if level == 3:
            p.paragraph_format.left_indent = Inches(0.32)
        elif level >= 4:
            p.paragraph_format.left_indent = Inches(0.58)
        apply_runs(p, parse_inline(text), spec_key="body")


# ---------------------------------------------------------------------------
# Document body clear
# ---------------------------------------------------------------------------

def _clear_body(doc: Document) -> None:
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def convert(md_path: Path, out_path: Path, template_path: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    blocks  = tokenize(md_text)
    toc_entries = _collect_static_toc_entries(blocks)
    doc     = Document(str(template_path))
    _clear_body(doc)

    ctx: str           = "body"
    front_matter: bool = False
    inserted_toc_heading = False
    saw_toc_heading = False
    skipping_source_toc = False

    def insert_toc_once() -> None:
        nonlocal inserted_toc_heading
        if inserted_toc_heading:
            return
        _para(doc, "body", "h2", "Table of Contents", force_bold=True)
        _insert_static_toc(doc, toc_entries)
        inserted_toc_heading = True

    for block in blocks:

        if block.kind == "heading":
            plain_heading = block.text.strip().lower()
            numbered_h1_section = block.level == 1 and _NUMBERED_SECTION_HEADING_RE.match(block.text.strip())
            if plain_heading == "table of contents":
                saw_toc_heading = True
                skipping_source_toc = True
                insert_toc_once()
                continue
            elif block.level >= 2 and not inserted_toc_heading:
                insert_toc_once()
            skipping_source_toc = False
            effective_level = 2 if numbered_h1_section else block.level
            style_key, spec_key = _HEADING_FORMAT.get(effective_level, ("body", "body"))
            new_ctx = _section_ctx(block.text)
            ctx = new_ctx if new_ctx else "body"
            if block.level == 1 and not numbered_h1_section:
                front_matter = True
            elif effective_level >= 2:
                front_matter = False
            _para(doc, style_key, spec_key, block.text)

        elif block.kind == "paragraph":
            text  = block.text.strip()
            plain = _plain_text(text)

            # Bold-only section label  e.g. **Abstract**
            new_ctx = _section_ctx(plain)
            if new_ctx:
                skipping_source_toc = False
                if new_ctx in {"abstract", "keywords"} and not inserted_toc_heading and not saw_toc_heading:
                    insert_toc_once()
                ctx = new_ctx
                _para(doc, "body", "body", text, force_bold=True)
                continue
            if skipping_source_toc:
                continue

            # Front matter: author / affiliation
            if front_matter and ctx == "body":
                if _AFFIL_START.match(text):
                    _para(doc, "address", "address", text)
                else:
                    _para(doc, "author", "author", text)
                continue

            if ctx != "body":
                spec = ctx if ctx in _FONT_SPEC else "body"
                _para(doc, ctx, spec, text)
            else:
                cap = _caption_style(text)
                key = cap if cap else "body"
                _para(doc, key, key, text)

        elif block.kind == "indented_block":
            if skipping_source_toc:
                continue
            for raw_line in block.lines:
                text = raw_line.strip()
                if not text:
                    continue
                cap = _caption_style(text)
                key = cap if cap else ("references" if ctx == "references" else "body")
                spec = key if key in _FONT_SPEC else "body"
                _para(doc, key, spec, text)

        elif block.kind == "ref_def":
            if skipping_source_toc:
                continue
            _para(doc, "references", "references", block.text)

        elif block.kind == "list_item":
            if skipping_source_toc:
                continue
            indent = "  " * block.depth
            bullet = (f"{indent}- {block.text}"
                      if not block.ordered else f"{indent}{block.text}")
            if ctx == "references":
                _para(doc, "references", "references", bullet)
            else:
                _para(doc, "body", "body", bullet)

        elif block.kind == "code_block":
            if skipping_source_toc:
                continue
            p  = doc.add_paragraph(style=_S["body"])
            wr = p.add_run(block.code)
            wr.font.name = "Courier New"
            wr.font.size = Pt(9)

        elif block.kind == "math_block":
            if skipping_source_toc:
                continue
            p = doc.add_paragraph(style=_S["body"])
            _apply_math(p, block.latex)

        elif block.kind == "table":
            if skipping_source_toc:
                continue
            _add_table(doc, block.header, block.rows)

        elif block.kind == "image":
            if skipping_source_toc:
                continue
            img_path = Path(block.path)
            if not img_path.is_absolute():
                img_path = md_path.parent / img_path
            if img_path.exists():
                p = doc.add_paragraph(style=_S["body"])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(img_path), width=Inches(_usable_page_width_inches(doc)))
                if block.alt:
                    _para(doc, "figure", "figure", f"Figure. {block.alt}")
            else:
                continue

        elif block.kind == "hr":
            # Horizontal rules in review Markdown are section separators, not
            # desired visual borders in the final DOCX.
            continue

    if not inserted_toc_heading and not saw_toc_heading:
        insert_toc_once()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    link_report = link_citations_and_dois(out_path)
    print(f"[md2docx] Saved -> {out_path}")
    print(
        "[md2docx] Linked citations/DOIs -> "
        f"{link_report['internal_hyperlink_count']} internal, "
        f"{link_report['doi_hyperlink_count']} DOI"
    )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

_DEFAULT_TEMPLATE = Path(__file__).resolve().parent.parent / "review_template.docx"


def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="md2docx",
        description="Convert Markdown to DOCX using review_template.docx styles.",
    )
    p.add_argument("--input",    required=True, metavar="MD",   help="Input .md file")
    p.add_argument("--output",   required=True, metavar="DOCX", help="Output .docx file")
    p.add_argument("--template", default=str(_DEFAULT_TEMPLATE), metavar="DOCX",
                   help=f"Word template (default: {_DEFAULT_TEMPLATE})")
    return p


def main() -> None:
    args          = _build_parser().parse_args()
    md_path       = Path(args.input).resolve()
    out_path      = Path(args.output).resolve()
    template_path = Path(args.template).resolve()

    if not md_path.exists():
        raise SystemExit(f"[md2docx] ERROR: Input not found: {md_path}")
    if not template_path.exists():
        raise SystemExit(f"[md2docx] ERROR: Template not found: {template_path}")
    if not _LATEX_OK:
        print("[md2docx] WARNING: latex2word not installed -- "
              "math will render as plain text.\n"
              "          Fix: pip install latex2word")

    convert(md_path, out_path, template_path)


if __name__ == "__main__":
    main()
